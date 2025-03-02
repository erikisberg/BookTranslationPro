import os
import PyPDF2
from fpdf import FPDF
import deepl
from openai import OpenAI
import tempfile
import logging
import time
import re
import io
from pathlib import Path
import subprocess
import shutil
import json

# Try to import optional document processing libraries
try:
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    # For the older pdfminer.six==20191110 used by textract
    from pdfminer.pdfinterp import PDFResourceManager, PDFPageInterpreter
    from pdfminer.converter import TextConverter
    from pdfminer.layout import LAParams
    from pdfminer.pdfpage import PDFPage
    from io import StringIO
    
    def pdfminer_extract_text(pdf_path):
        """Extract text from PDF using older pdfminer.six API"""
        resource_manager = PDFResourceManager()
        fake_file_handle = StringIO()
        converter = TextConverter(resource_manager, fake_file_handle, laparams=LAParams())
        page_interpreter = PDFPageInterpreter(resource_manager, converter)
        
        with open(pdf_path, 'rb') as fh:
            for page in PDFPage.get_pages(fh, caching=True, check_extractable=True):
                page_interpreter.process_page(page)
                
            text = fake_file_handle.getvalue()
            
        converter.close()
        fake_file_handle.close()
        
        return text
        
    PDFMINER_AVAILABLE = True
except ImportError:
    PDFMINER_AVAILABLE = False

try:
    import textract
    TEXTRACT_AVAILABLE = True
except ImportError:
    TEXTRACT_AVAILABLE = False

logger = logging.getLogger(__name__)

def is_allowed_file(filename):
    """Check if the file type is supported for processing"""
    if not filename:
        return False
    
    allowed_extensions = [
        '.pdf',  # Adobe PDF
        '.docx', '.doc',  # Microsoft Word
        '.txt',  # Plain text
        '.rtf',  # Rich Text Format
        '.odt'   # OpenDocument Text
    ]
    
    return any(filename.lower().endswith(ext) for ext in allowed_extensions)

def extract_text_from_pdf(filepath):
    """Extract text from a PDF file using multiple methods for reliability"""
    pages_text = []
    fallback_used = False
    
    # First try with PyPDF2
    try:
        with open(filepath, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            total_pages = len(pdf_reader.pages)
            logger.info(f"Processing PDF with {total_pages} pages using PyPDF2")
            
            for page_num in range(total_pages):
                try:
                    page = pdf_reader.pages[page_num]
                    text = page.extract_text()
                    
                    # Check if text was extracted
                    if text and not text.isspace() and len(text.strip()) > 10:
                        pages_text.append({
                            'id': page_num,
                            'text': text.strip(),
                            'source': 'PyPDF2'
                        })
                        logger.info(f"Successfully extracted {len(text)} characters from page {page_num + 1} with PyPDF2")
                    else:
                        # If PyPDF2 fails, flag for fallback methods
                        fallback_used = True
                        logger.warning(f"PyPDF2 extracted insufficient text from page {page_num + 1}, will try fallback")
                
                except Exception as e:
                    fallback_used = True
                    logger.error(f"PyPDF2 error on page {page_num + 1}: {str(e)}")
    
    except Exception as e:
        fallback_used = True
        logger.error(f"PyPDF2 failed to process PDF: {str(e)}")
    
    # If PyPDF2 couldn't extract text from some pages, try pdfminer
    if fallback_used and PDFMINER_AVAILABLE:
        try:
            logger.info(f"Trying PDFMiner for better text extraction")
            full_text = pdfminer_extract_text(filepath)
            
            # If we got text but haven't processed any pages successfully yet
            if full_text and len(pages_text) == 0:
                # Split into pages based on common page break markers
                page_texts = re.split(r'\f|\n\s*\n\s*\n', full_text)
                
                for i, page_text in enumerate(page_texts):
                    if page_text and not page_text.isspace() and len(page_text.strip()) > 10:
                        pages_text.append({
                            'id': i,
                            'text': page_text.strip(),
                            'source': 'PDFMiner'
                        })
                        logger.info(f"Successfully extracted text from section {i+1} with PDFMiner")
            
            # If we have page numbers with missing text, fill them in
            elif len(pages_text) > 0:
                # Get page numbers that failed with PyPDF2
                failed_pages = []
                extracted_ids = [page['id'] for page in pages_text]
                
                with open(filepath, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    for i in range(len(pdf_reader.pages)):
                        if i not in extracted_ids:
                            failed_pages.append(i)
                
                # Try to estimate page boundaries and insert content
                if failed_pages and full_text:
                    logger.info(f"Attempting to fix {len(failed_pages)} pages with PDFMiner")
                    # This is a simplistic approach - in a real app, you'd need more sophisticated page detection
                    approx_page_length = len(full_text) / len(pdf_reader.pages)
                    
                    for page_num in failed_pages:
                        start_pos = int(page_num * approx_page_length)
                        end_pos = int((page_num + 1) * approx_page_length)
                        if start_pos < len(full_text) and end_pos <= len(full_text):
                            page_text = full_text[start_pos:end_pos].strip()
                            if len(page_text) > 10:
                                pages_text.append({
                                    'id': page_num,
                                    'text': page_text,
                                    'source': 'PDFMiner-estimated'
                                })
                                logger.info(f"Added estimated text for page {page_num+1} with PDFMiner")
        
        except Exception as e:
            logger.error(f"PDFMiner fallback failed: {str(e)}")
    
    # Last resort - try textract if available
    if len(pages_text) == 0 and TEXTRACT_AVAILABLE:
        try:
            logger.info("Trying textract as last resort")
            text = textract.process(filepath).decode('utf-8')
            if text and len(text.strip()) > 10:
                # We don't know page boundaries, so just create a single "page"
                pages_text.append({
                    'id': 0,
                    'text': text.strip(),
                    'source': 'textract'
                })
                logger.info(f"Extracted {len(text)} characters with textract")
        except Exception as e:
            logger.error(f"Textract fallback failed: {str(e)}")
    
    # Sort by page number
    pages_text.sort(key=lambda x: x['id'])
    
    if not pages_text:
        raise Exception("Could not extract any text from the PDF with any method")
    
    logger.info(f"Successfully extracted text from {len(pages_text)} pages/sections")
    return pages_text

def extract_text_from_docx(filepath):
    """Extract text from a DOCX file"""
    if not DOCX_AVAILABLE:
        raise ImportError("python-docx package is not installed")
    
    try:
        doc = docx.Document(filepath)
        sections = []
        
        # Option 1: Extract by paragraphs
        current_section = []
        section_count = 0
        
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if text:
                current_section.append(text)
            
            # Create a new section after a certain number of paragraphs
            # or if we encounter a page break (simplistic detection)
            if (i > 0 and i % 10 == 0) or (text.startswith('\f') or not text and len(current_section) > 5):
                section_text = '\n'.join(current_section)
                if len(section_text) > 10:
                    sections.append({
                        'id': section_count,
                        'text': section_text,
                        'source': 'docx-paragraphs'
                    })
                    section_count += 1
                    current_section = []
        
        # Add the final section if it exists
        if current_section:
            section_text = '\n'.join(current_section)
            if len(section_text) > 10:
                sections.append({
                    'id': section_count,
                    'text': section_text,
                    'source': 'docx-paragraphs'
                })
        
        # If we couldn't parse meaningful sections, try a simpler approach
        if not sections:
            full_text = '\n'.join([p.text for p in doc.paragraphs if p.text.strip()])
            if full_text and len(full_text) > 10:
                sections.append({
                    'id': 0,
                    'text': full_text,
                    'source': 'docx-full'
                })
        
        if not sections:
            raise Exception("Could not extract meaningful text from DOCX")
            
        logger.info(f"Successfully extracted {len(sections)} sections from DOCX")
        return sections
        
    except Exception as e:
        logger.error(f"Error extracting text from DOCX: {str(e)}")
        raise

def extract_text_from_txt(filepath):
    """Extract text from a plain text file"""
    try:
        with open(filepath, 'r', encoding='utf-8', errors='replace') as file:
            text = file.read()
        
        # Split into manageable sections (e.g., by double line breaks)
        sections = re.split(r'\n\s*\n', text)
        pages_text = []
        
        # If we have very few sections, try to split by other means
        if len(sections) <= 1:
            # Try splitting by regular expressions matching potential section headers
            sections = re.split(r'\n\s*(?:[A-Z][A-Z\s]+:|\d+\.\s+[A-Z]|\*\*\*|\-{3,}|\={3,})', text)
        
        # If we still have only one section, create artificial sections
        if len(sections) <= 1 and len(text) > 5000:
            # Create sections of about 3000 characters each, at line breaks
            lines = text.split('\n')
            section = []
            char_count = 0
            section_count = 0
            
            for line in lines:
                section.append(line)
                char_count += len(line)
                
                if char_count > 3000 and line.strip():
                    section_text = '\n'.join(section)
                    pages_text.append({
                        'id': section_count,
                        'text': section_text,
                        'source': 'txt-chunked'
                    })
                    section_count += 1
                    section = []
                    char_count = 0
            
            # Add the final section
            if section:
                section_text = '\n'.join(section)
                pages_text.append({
                    'id': section_count,
                    'text': section_text,
                    'source': 'txt-chunked'
                })
        else:
            # Use the natural sections we found
            for i, section_text in enumerate(sections):
                if section_text.strip() and len(section_text.strip()) > 10:
                    pages_text.append({
                        'id': i,
                        'text': section_text.strip(),
                        'source': 'txt-sections'
                    })
        
        if not pages_text:
            # Just use the whole text as one section
            pages_text.append({
                'id': 0,
                'text': text.strip(),
                'source': 'txt-full'
            })
        
        logger.info(f"Successfully extracted {len(pages_text)} sections from text file")
        return pages_text
        
    except Exception as e:
        logger.error(f"Error extracting text from text file: {str(e)}")
        raise

def extract_text_from_file(filepath):
    """Extract text from various file formats"""
    file_ext = os.path.splitext(filepath)[1].lower()
    
    if file_ext in ['.pdf']:
        return extract_text_from_pdf(filepath)
    elif file_ext in ['.docx', '.doc']:
        # For .doc, try to convert to .docx first if possible
        if file_ext == '.doc' and shutil.which('soffice'):
            try:
                logger.info("Converting .doc to .docx for better text extraction")
                temp_dir = tempfile.mkdtemp()
                output_path = os.path.join(temp_dir, 'converted.docx')
                
                # Use LibreOffice to convert
                subprocess.run([
                    'soffice', '--headless', '--convert-to', 'docx', 
                    '--outdir', temp_dir, filepath
                ], check=True, capture_output=True)
                
                if os.path.exists(output_path):
                    result = extract_text_from_docx(output_path)
                    shutil.rmtree(temp_dir, ignore_errors=True)
                    return result
            except Exception as e:
                logger.error(f"Error converting .doc to .docx: {str(e)}")
                # Continue with original file
        
        # If conversion failed or wasn't attempted, try direct extraction
        if DOCX_AVAILABLE:
            try:
                return extract_text_from_docx(filepath)
            except Exception:
                pass
                
        # Fallback to textract if available
        if TEXTRACT_AVAILABLE:
            try:
                text = textract.process(filepath).decode('utf-8')
                if text and len(text.strip()) > 10:
                    return [{
                        'id': 0,
                        'text': text.strip(),
                        'source': 'textract'
                    }]
            except Exception as e:
                logger.error(f"Textract fallback failed for DOCX: {str(e)}")
        
        raise Exception(f"Could not extract text from {file_ext} file")
        
    elif file_ext in ['.txt', '.rtf']:
        # For RTF, try to convert to txt first if possible
        if file_ext == '.rtf' and TEXTRACT_AVAILABLE:
            try:
                text = textract.process(filepath).decode('utf-8')
                return [{
                    'id': 0,
                    'text': text.strip(),
                    'source': 'textract-rtf'
                }]
            except:
                pass
        
        # Plain text
        return extract_text_from_txt(filepath)
    
    elif file_ext in ['.odt']:
        # Use textract if available
        if TEXTRACT_AVAILABLE:
            try:
                text = textract.process(filepath).decode('utf-8')
                if text and len(text.strip()) > 10:
                    return [{
                        'id': 0,
                        'text': text.strip(),
                        'source': 'textract-odt'
                    }]
            except Exception as e:
                logger.error(f"Textract failed for ODT: {str(e)}")
        
        # Try LibreOffice conversion as fallback
        if shutil.which('soffice'):
            try:
                temp_dir = tempfile.mkdtemp()
                output_path = os.path.join(temp_dir, 'converted.txt')
                
                # Use LibreOffice to convert
                subprocess.run([
                    'soffice', '--headless', '--convert-to', 'txt:Text', 
                    '--outdir', temp_dir, filepath
                ], check=True, capture_output=True)
                
                if os.path.exists(output_path):
                    result = extract_text_from_txt(output_path)
                    shutil.rmtree(temp_dir, ignore_errors=True)
                    return result
            except Exception as e:
                logger.error(f"Error converting ODT: {str(e)}")
        
        raise Exception("Could not extract text from ODT file")
    
    else:
        raise ValueError(f"Unsupported file format: {file_ext}")

def extract_text_from_page(pdf_reader, page_num):
    """Legacy function for backward compatibility"""
    try:
        page = pdf_reader.pages[page_num]
        text = page.extract_text()

        # Validate extracted text
        if not text or text.isspace():
            raise ValueError("Extracted text is empty or whitespace only")

        text = text.strip()
        if len(text) < 10:  # Minimum reasonable length for a page
            raise ValueError("Extracted text is too short to be valid content")

        logger.info(f"Successfully extracted {len(text)} characters from page {page_num + 1}")
        return text

    except Exception as e:
        logger.error(f"Error extracting text from page {page_num + 1}: {str(e)}")
        raise Exception(f"Failed to extract text from page {page_num + 1}: {str(e)}")

def translate_text(text, deepl_api_key, target_language='SV', source_language='auto', use_cache=True, glossary_id=None, max_retries=3, timeout=30):
    """First step: Translate text using DeepL with caching and glossary support.
    
    Returns a tuple containing (translated_text, text_hash, source_text, glossary_hits, glossary_terms_used).
    
    Args:
        text: The text to translate
        deepl_api_key: DeepL API key
        target_language: Target language code (default: 'SV' for Swedish)
        source_language: Source language code or 'auto' for auto-detection
        use_cache: Whether to use translation memory cache
        glossary_id: ID of glossary to apply (optional)
        max_retries: Maximum number of retries for API failures
        timeout: Timeout in seconds for API calls
        
    Returns:
        Tuple containing (translated_text, text_hash, source_text, glossary_hits, glossary_terms_used)
        
    Raises:
        ValueError: If input is empty or API key is invalid
        Exception: For other translation errors
    """
    # Validate input
    if not text or text.isspace():
        logger.warning("Received empty text for translation")
        raise ValueError("Cannot translate empty text")
    
    # Validate API key format (basic check)
    if not deepl_api_key or len(deepl_api_key) < 20:
        logger.error("Invalid DeepL API key format")
        raise ValueError("DeepL API key appears to be invalid (too short)")
    
    # Validate language codes
    VALID_TARGET_LANGUAGES = ['BG', 'CS', 'DA', 'DE', 'EL', 'EN', 'ES', 'ET', 'FI', 'FR', 'HU', 'ID', 'IT', 'JA', 'KO', 'LT', 'LV', 'NB', 'NL', 'PL', 'PT', 'RO', 'RU', 'SK', 'SL', 'SV', 'TR', 'UK', 'ZH']
    VALID_SOURCE_LANGUAGES = ['AUTO', 'BG', 'CS', 'DA', 'DE', 'EL', 'EN', 'ES', 'ET', 'FI', 'FR', 'HU', 'ID', 'IT', 'JA', 'KO', 'LT', 'LV', 'NB', 'NL', 'PL', 'PT', 'RO', 'RU', 'SK', 'SL', 'SV', 'TR', 'UK', 'ZH']
    
    # Normalize language codes to uppercase
    target_language = target_language.upper()
    source_language = source_language.upper()
    
    if target_language not in VALID_TARGET_LANGUAGES:
        logger.warning(f"Invalid target language code: {target_language}, using EN")
        target_language = 'EN'  # Fallback to English
    
    if source_language not in VALID_SOURCE_LANGUAGES:
        logger.warning(f"Invalid source language code: {source_language}, using AUTO")
        source_language = 'AUTO'  # Fallback to auto-detection
        
    # Critical check: if source and target languages are the same, force source to AUTO
    # This prevents issues where both source and target become the same text
    if source_language == target_language and source_language != 'AUTO':
        logger.warning(f"Source and target languages are the same ({source_language}). Forcing source to AUTO")
        source_language = 'AUTO'

    # Default values for glossary stats
    glossary_hits = 0
    glossary_terms_used = 0

    # Check cache first if enabled
    if use_cache:
        try:
            from supabase_config import check_translation_cache, generate_text_hash
            
            # Look for an existing translation in cache
            cached_translation = check_translation_cache(text, target_language)
            if cached_translation:
                logger.info("Translation found in cache, skipping DeepL API call")
                
                # Apply glossary to cached translation if needed
                if glossary_id:
                    try:
                        from supabase_config import apply_glossary_to_text
                        glossary_result = apply_glossary_to_text(cached_translation, glossary_id)
                        
                        # Unpack the result tuple (modified_text, replacements_count, entry_count)
                        if isinstance(glossary_result, tuple) and len(glossary_result) == 3:
                            glossary_applied_text, glossary_hits, glossary_terms_used = glossary_result
                            logger.info(f"Applied glossary to cached translation ({glossary_hits} replacements from {glossary_terms_used} terms)")
                        else:
                            # Fallback for older version of apply_glossary_to_text
                            glossary_applied_text = glossary_result
                            logger.info(f"Applied glossary to cached translation with older function version")
                        
                        return (
                            glossary_applied_text, 
                            generate_text_hash(text, target_language), 
                            text, 
                            glossary_hits, 
                            glossary_terms_used
                        )
                    except Exception as glossary_error:
                        logger.error(f"Error applying glossary to cached translation: {str(glossary_error)}")
                        # Fall back to cached translation without glossary
                
                return cached_translation, generate_text_hash(text, target_language), text, 0, 0
        except Exception as cache_error:
            # If cache check fails, log but continue with normal translation
            logger.error(f"Error checking translation cache: {str(cache_error)}")

    # No cache hit, use DeepL with retry mechanism
    logger.info(f"DeepL API Key starting with: {deepl_api_key[:5]}...")
    logger.info(f"Source language: {source_language}, Target language: {target_language}")
    
    # Define retry mechanism with exponential backoff
    retry_count = 0
    last_error = None
    
    while retry_count < max_retries:
        try:
            # Create translator with timeout
            translator = deepl.Translator(deepl_api_key, timeout=timeout)
            
            # Validate API key by checking usage
            try:
                usage = translator.get_usage()
                logger.debug(f"DeepL API usage: {usage.character.count}/{usage.character.limit} characters")
                
                # Check if we're near the limit
                if usage.character.limit > 0 and usage.character.count / usage.character.limit > 0.95:
                    logger.warning(f"DeepL API usage is at {usage.character.count}/{usage.character.limit} characters (95%+ of limit)")
            except Exception as usage_error:
                logger.warning(f"Could not check DeepL API usage: {str(usage_error)}")
            
            logger.info(f"Sending {len(text)} characters to DeepL for translation")
            logger.info(f"Text sample: {text[:100]}...")
            
            # Use source_language only if it's not auto-detect
            if source_language == 'AUTO':
                result = translator.translate_text(text, target_lang=target_language)
            else:
                result = translator.translate_text(text, source_lang=source_language, target_lang=target_language)

            # Validate response
            if not result:
                raise ValueError("DeepL returned None result")
                
            if not hasattr(result, 'text'):
                raise ValueError("DeepL returned malformed result without text attribute")
                
            if not result.text or result.text.isspace():
                raise ValueError("DeepL returned empty translation")

            logger.info("Text successfully translated with DeepL")
            logger.info(f"Translation sample: {result.text[:100]}...")
            
            # Apply glossary to translation if specified
            translation_text = result.text
            if glossary_id:
                try:
                    from supabase_config import apply_glossary_to_text
                    glossary_result = apply_glossary_to_text(translation_text, glossary_id)
                    
                    # Unpack the result tuple (modified_text, replacements_count, entry_count)
                    if isinstance(glossary_result, tuple) and len(glossary_result) == 3:
                        translation_text, glossary_hits, glossary_terms_used = glossary_result
                        logger.info(f"Applied glossary to DeepL translation ({glossary_hits} replacements from {glossary_terms_used} terms)")
                    else:
                        # Fallback for older version of apply_glossary_to_text
                        translation_text = glossary_result
                        logger.info(f"Applied glossary to DeepL translation with older function version")
                except Exception as glossary_error:
                    logger.error(f"Error applying glossary to translation: {str(glossary_error)}")
                    # Continue with the original translation
            
            # Generate hash for caching if needed
            text_hash = None
            if use_cache:
                try:
                    from supabase_config import generate_text_hash
                    text_hash = generate_text_hash(text, target_language)
                except Exception as hash_error:
                    logger.error(f"Error generating hash for caching: {str(hash_error)}")
            
            return translation_text, text_hash, text, glossary_hits, glossary_terms_used

        except deepl.exceptions.AuthorizationException as auth_err:
            # Authentication errors won't be fixed by retrying
            logger.error(f"DeepL API authentication error: {str(auth_err)}")
            raise ValueError(f"DeepL API key is invalid or has expired: {str(auth_err)}")
            
        except deepl.exceptions.QuotaExceededException as quota_err:
            # Quota errors won't be fixed by retrying
            logger.error(f"DeepL API quota exceeded: {str(quota_err)}")
            raise ValueError(f"DeepL API quota has been exceeded: {str(quota_err)}")
            
        except (deepl.exceptions.ConnectionException, deepl.exceptions.DeepLException) as network_err:
            # Network/transient errors are worth retrying with backoff
            retry_count += 1
            wait_time = min(2 ** retry_count, 60)  # Exponential backoff capped at 60 seconds
            
            logger.warning(f"DeepL API error (attempt {retry_count}/{max_retries}): {str(network_err)}. Retrying in {wait_time} seconds...")
            last_error = network_err
            
            # Wait before retrying
            time.sleep(wait_time)
            
        except Exception as e:
            # Other unexpected errors
            logger.error(f"Unexpected DeepL translation error: {str(e)}")
            raise Exception(f"Translation failed: {str(e)}")
    
    # If we've exhausted retries
    if last_error:
        logger.error(f"DeepL translation failed after {max_retries} attempts: {str(last_error)}")
        raise Exception(f"Translation failed after {max_retries} attempts: {str(last_error)}")
        
    # Fallback - we should never reach here, but just in case
    raise Exception("Translation failed: Unknown error")

def create_openai_assistant(openai_api_key, name, instructions, model="gpt-4o"):
    """Create a new OpenAI assistant with the given name and instructions"""
    try:
        logger.info(f"Creating new OpenAI assistant: {name}")
        client = OpenAI(api_key=openai_api_key)
        
        # Create the assistant
        assistant = client.beta.assistants.create(
            name=name,
            instructions=instructions,
            model=model,
            tools=[]  # No tools needed for translation review
        )
        
        logger.info(f"Successfully created assistant with ID: {assistant.id}")
        return {
            "id": assistant.id,
            "name": assistant.name,
            "instructions": assistant.instructions,
            "model": assistant.model
        }
    except Exception as e:
        logger.error(f"Error creating OpenAI assistant: {str(e)}")
        raise Exception(f"Failed to create OpenAI assistant: {str(e)}")

def update_openai_assistant(openai_api_key, assistant_id, name=None, instructions=None, model="gpt-4o"):
    """Update an existing OpenAI assistant with new name, instructions, or model"""
    try:
        logger.info(f"Updating OpenAI assistant: {assistant_id}")
        client = OpenAI(api_key=openai_api_key)
        
        # Get current assistant data to update only what's provided
        current = client.beta.assistants.retrieve(assistant_id)
        
        # Prepare update parameters
        update_params = {}
        if name is not None:
            update_params["name"] = name
        if instructions is not None:
            update_params["instructions"] = instructions
        if model is not None:
            update_params["model"] = model
        
        # Only update if there are changes
        if update_params:
            assistant = client.beta.assistants.update(
                assistant_id=assistant_id,
                **update_params
            )
            
            logger.info(f"Successfully updated assistant: {assistant.id}")
            return {
                "id": assistant.id,
                "name": assistant.name,
                "instructions": assistant.instructions,
                "model": assistant.model
            }
        else:
            logger.info(f"No changes to update for assistant: {assistant_id}")
            return {
                "id": current.id,
                "name": current.name,
                "instructions": current.instructions,
                "model": current.model
            }
    except Exception as e:
        logger.error(f"Error updating OpenAI assistant: {str(e)}")
        raise Exception(f"Failed to update OpenAI assistant: {str(e)}")

def delete_openai_assistant(openai_api_key, assistant_id):
    """Delete an OpenAI assistant"""
    try:
        logger.info(f"Deleting OpenAI assistant: {assistant_id}")
        client = OpenAI(api_key=openai_api_key)
        
        deletion = client.beta.assistants.delete(assistant_id)
        logger.info(f"Assistant deletion response: {deletion}")
        
        return deletion.deleted
    except Exception as e:
        logger.error(f"Error deleting OpenAI assistant: {str(e)}")
        raise Exception(f"Failed to delete OpenAI assistant: {str(e)}")

def review_translation(text, openai_api_key, assistant_id, instructions=None, max_retries=3, timeout=300):
    """Second step: Review the translation using OpenAI Assistant with robust error handling.
    
    Args:
        text: The text to review
        openai_api_key: OpenAI API key
        assistant_id: ID of the assistant to use
        instructions: Custom instructions for the review (optional)
        max_retries: Maximum number of retry attempts for transient errors
        timeout: Maximum time to wait for completion in seconds
        
    Returns:
        The reviewed translation text, or the original text if review fails
        
    Raises:
        ValueError: For validation errors
        Exception: For critical errors (optional, can be caught internally)
    """
    # Validate inputs
    if not text or text.isspace():
        logger.warning("Cannot review empty translation")
        raise ValueError("Cannot review empty translation")
        
    # Validate API key format (basic check)
    if not openai_api_key or len(openai_api_key) < 30:
        logger.error("Invalid OpenAI API key format")
        raise ValueError("OpenAI API key appears to be invalid (too short)")
        
    # Validate assistant ID format
    if not assistant_id or not isinstance(assistant_id, str) or len(assistant_id) < 10:
        logger.error(f"Invalid assistant ID format: {assistant_id}")
        raise ValueError("Assistant ID appears to be invalid")
        
    # Limit text length to avoid excessive costs
    MAX_CHARS = 100000  # 100k character limit
    if len(text) > MAX_CHARS:
        logger.warning(f"Text too long for review ({len(text)} chars). Truncating to {MAX_CHARS} chars.")
        text = text[:MAX_CHARS]

    logger.info(f"OpenAI API Key starting with: {openai_api_key[:5]}...")
    logger.info(f"Assistant ID: {assistant_id}")
    
    # Use the specified OpenAI API key
    client = OpenAI(api_key=openai_api_key)
    
    # Track start time for timeout
    start_time = time.time()
    
    # Create a thread ID outside the try block for potential cleanup
    thread_id = None

    # Define retry mechanism with backoff
    retry_count = 0
    last_error = None
    
    while retry_count < max_retries:
        try:
            # Check if we've exceeded the timeout
            elapsed_time = time.time() - start_time
            if elapsed_time > timeout:
                logger.error(f"Review timed out after {elapsed_time:.1f} seconds")
                return text
            
            # First validate the assistant exists
            try:
                # Verify this is even a valid assistant ID format
                if not assistant_id or not str(assistant_id).startswith("asst_"):
                    logger.error(f"Invalid OpenAI assistant ID format: {assistant_id}. Must start with 'asst_'")
                    return text  # Return original text without review

                # Retrieve the assistant
                assistant = client.beta.assistants.retrieve(assistant_id)
                logger.info(f"Using assistant: {assistant.name} (model: {assistant.model})")
            except Exception as assistant_error:
                logger.error(f"Could not retrieve assistant with ID {assistant_id}: {assistant_error}")
                # Return original text instead of raising error
                return text
            
            # Create a new thread
            logger.info(f"Creating new thread for reviewing {len(text)} characters")
            thread = client.beta.threads.create()
            thread_id = thread.id
            logger.info(f"Created thread with ID: {thread_id}")

            # Use custom instructions if provided, otherwise use default
            if not instructions:
                instructions = """Granska och förbättra denna svenska översättning. 
                Texten ska vara tydlig, naturlig och bevara den ursprungliga betydelsen.
                Om du inte kan förbättra texten, returnera den som den är."""
                
            # Sanitize the instructions and text
            if instructions and len(instructions) > 5000:
                logger.warning(f"Instructions too long ({len(instructions)} chars). Truncating.")
                instructions = instructions[:5000]
                
            # Construct message with clear instructions
            logger.info("Adding message to thread")
            message_content = f"""{instructions}

                ---

                {text}

                ---

                If the text seems confused or unclear, return an error message starting with 'TRANSLATION_ERROR:'.
                """
            logger.info(f"Message sample (first 100 chars): {message_content[:100]}...")
            
            # Create the message in the thread
            message = client.beta.threads.messages.create(
                thread_id=thread_id,
                role="user",
                content=message_content
            )
            
            if not message or not message.id:
                raise ValueError("Failed to create message in thread")

            # Start the assistant run
            logger.info("Starting assistant run")
            run = client.beta.threads.runs.create(
                thread_id=thread_id,
                assistant_id=assistant_id
            )
            
            if not run or not run.id:
                raise ValueError("Failed to start assistant run")
                
            run_id = run.id
            logger.info(f"Started run with ID: {run_id}")

            # Poll for completion with improved exponential backoff
            poll_count = 0
            max_polls = 30  # Safety limit to avoid infinite loops
            current_delay = 5  # Start with a 5 second delay
            max_poll_delay = 30  # Maximum delay between polls
            
            while poll_count < max_polls:
                # Check if we've exceeded the timeout
                elapsed_time = time.time() - start_time
                if elapsed_time > timeout:
                    logger.error(f"Review timed out after {elapsed_time:.1f} seconds")
                    return text
                
                # Wait before polling
                if poll_count > 0:
                    logger.info(f"Waiting {current_delay} seconds before next check")
                    time.sleep(current_delay)
                    # Increase delay for next poll with cap
                    current_delay = min(current_delay * 1.5, max_poll_delay)
                
                # Poll for status
                logger.info(f"Checking run status (poll {poll_count + 1}/{max_polls}, elapsed: {elapsed_time:.1f}s)")
                try:
                    run_status = client.beta.threads.runs.retrieve(
                        thread_id=thread_id,
                        run_id=run_id
                    )
                    
                    if not run_status:
                        logger.warning("Empty run status response")
                        poll_count += 1
                        continue
                except Exception as poll_error:
                    logger.warning(f"Error polling run status: {poll_error}")
                    poll_count += 1
                    continue
                
                # Check status
                status = run_status.status
                logger.info(f"Run status: {status}")
                
                if status == 'completed':
                    logger.info(f"Run completed successfully after {elapsed_time:.1f} seconds")
                    break
                elif status in ['failed', 'cancelled', 'expired']:
                    logger.error(f"Run failed with status: {status}")
                    # Check for failure details
                    try:
                        if hasattr(run_status, 'last_error') and run_status.last_error:
                            logger.error(f"Error details: {run_status.last_error}")
                    except:
                        pass
                    return text
                elif status in ['queued', 'in_progress', 'requires_action']:
                    # Continue polling
                    poll_count += 1
                    continue
                else:
                    # Unknown status
                    logger.warning(f"Unknown run status: {status}")
                    poll_count += 1
            
            # If we exited the loop without completing, return original text
            if status != 'completed':
                logger.warning(f"Run did not complete after {max_polls} polls. Status: {status}")
                return text

            # Successfully completed, retrieve messages
            logger.info("Retrieving assistant's response")
            try:
                messages = client.beta.threads.messages.list(thread_id=thread_id)
            except Exception as msg_error:
                logger.error(f"Error retrieving messages: {msg_error}")
                return text

            if not messages or not hasattr(messages, 'data') or not messages.data:
                logger.warning("No messages received from assistant")
                return text

            # Find the assistant's response (should be the newest message from assistant)
            assistant_messages = [msg for msg in messages.data 
                                 if msg.role == 'assistant' and msg.content]
            
            if not assistant_messages:
                logger.warning("No assistant messages found in response")
                return text
                
            # Get the newest message
            assistant_message = assistant_messages[0]
            
            try:
                # Extract text content from the message
                text_contents = [content.text.value 
                               for content in assistant_message.content 
                               if hasattr(content, 'text') and hasattr(content.text, 'value')]
                
                if not text_contents:
                    logger.warning("No text content found in assistant message")
                    return text
                
                # Join all text content fragments
                response = '\n'.join(text_contents)
                
                # Check for error message pattern
                if response.startswith('TRANSLATION_ERROR:'):
                    logger.warning(f"Assistant reported translation error: {response}")
                    return text  # Return original on error

                logger.info("Translation review completed successfully")
                logger.info(f"Response sample (first 100 chars): {response[:100]}...")
                
                # Clean up the thread to avoid cluttering OpenAI storage
                try:
                    client.beta.threads.delete(thread_id)
                    logger.debug(f"Deleted thread {thread_id}")
                except Exception as cleanup_error:
                    logger.warning(f"Error cleaning up thread: {cleanup_error}")
                
                return response  # Success!

            except Exception as content_error:
                logger.error(f"Error extracting message content: {content_error}")
                return text

        except openai.AuthenticationError as auth_err:
            logger.error(f"OpenAI API authentication error: {auth_err}")
            # Don't retry auth errors
            return text
            
        except openai.RateLimitError as rate_err:
            # Rate limiting errors might be retryable with longer delays
            retry_count += 1
            last_error = rate_err
            wait_time = min(30 * retry_count, 120)  # 30s, 60s, 90s, capped at 120s
            logger.warning(f"OpenAI API rate limited (attempt {retry_count}/{max_retries}). Retrying in {wait_time}s...")
            time.sleep(wait_time)
            
        except openai.APITimeoutError as timeout_err:
            # Timeouts are retryable
            retry_count += 1
            last_error = timeout_err
            wait_time = min(10 * retry_count, 60)
            logger.warning(f"OpenAI API timeout (attempt {retry_count}/{max_retries}). Retrying in {wait_time}s...")
            time.sleep(wait_time)
            
        except (openai.APIConnectionError, openai.BadRequestError) as conn_err:
            # Connection errors can be retried
            retry_count += 1
            last_error = conn_err
            wait_time = min(5 * retry_count, 30)
            logger.warning(f"OpenAI API connection error (attempt {retry_count}/{max_retries}): {conn_err}. Retrying in {wait_time}s...")
            time.sleep(wait_time)
            
        except Exception as e:
            # For other errors, log and return original text
            logger.error(f"Unexpected error in review_translation: {e}")
            # Clean up thread if it was created
            if thread_id:
                try:
                    client.beta.threads.delete(thread_id)
                    logger.debug(f"Deleted thread {thread_id} after error")
                except:
                    pass
            return text
            
    # If we've exhausted retries
    if last_error:
        logger.error(f"Review failed after {max_retries} attempts: {last_error}")
        
    # Clean up thread if it was created
    if thread_id:
        try:
            client.beta.threads.delete(thread_id)
            logger.debug(f"Deleted thread {thread_id} after exhausting retries")
        except:
            pass
            
    return text  # Return original text if all retries failed

def create_pdf_with_text(text_content):
    """Legacy function - kept for backward compatibility"""
    return create_pdf_with_formatting(text_content)
    
def create_pdf_with_text_basic(text_content):
    """Simple PDF creation as a fallback method"""
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font('helvetica', size=12)
    pdf.set_auto_page_break(auto=True, margin=15)

    try:
        paragraphs = text_content.split('\n')
        for paragraph in paragraphs:
            cleaned_text = paragraph.encode('latin-1', errors='replace').decode('latin-1')
            pdf.multi_cell(0, 10, cleaned_text)
            pdf.ln(5)
    except Exception as e:
        logger.error(f"Error in basic PDF creation: {str(e)}")
        # If even this fails, just create an empty PDF
        pass

    temp_output = tempfile.NamedTemporaryFile(delete=False, suffix='.pdf')
    pdf.output(temp_output.name)
    return temp_output.name

def create_pdf_with_formatting(text_content, font_family='helvetica', font_size=12, page_size='A4', 
                               orientation='portrait', margin=15, line_spacing=1.5, alignment='left',
                               include_page_numbers=True, header_text='', footer_text=''):
    """Create a PDF with the given text content using the specified formatting options"""
    try:
        logger.info(f"Creating PDF with settings: font={font_family}, size={font_size}, page={page_size}, orientation={orientation}")
        
        # Create a custom PDF class with header and footer methods to avoid recursion
        class CustomPDF(FPDF):
            def header(self):
                if header_text:
                    self.set_y(5)
                    self.set_font(font_family, 'I', font_size - 2)
                    self.cell(0, 10, header_text, 0, 1, 'C')
                    self.set_y(margin)
                    self.set_font(font_family, size=font_size)
                    
            def footer(self):
                if footer_text or include_page_numbers:
                    self.set_y(-15)
                    self.set_font(font_family, 'I', font_size - 2)
                    
                    if footer_text and include_page_numbers:
                        footer = f"{footer_text} | Sida {self.page_no()}"
                    elif footer_text:
                        footer = footer_text
                    elif include_page_numbers:
                        footer = f"Sida {self.page_no()}"
                    else:
                        return
                        
                    self.cell(0, 10, footer, 0, 0, 'C')
        
        pdf = CustomPDF(orientation=orientation[0], format=page_size)
        pdf.add_page()
        pdf.set_font(font_family, size=font_size)
        pdf.set_auto_page_break(auto=True, margin=margin)
        
        # Set margins
        pdf.set_margins(margin, margin, margin)
        
        # Footer is now handled by the CustomPDF class
        
        # Set alignment
        align_dict = {
            'left': 'L',
            'center': 'C', 
            'right': 'R',
            'justified': 'J'
        }
        align = align_dict.get(alignment, 'L')
        
        # Calculate line height based on line spacing
        line_height = font_size * 0.5 * line_spacing
        
        paragraphs = text_content.split('\n')
        for paragraph in paragraphs:
            if not paragraph.strip():
                pdf.ln(line_height)
                continue
                
            # Handle text encoding
            cleaned_text = paragraph.encode('latin-1', errors='replace').decode('latin-1')
            pdf.multi_cell(0, line_height, cleaned_text, 0, align)
            pdf.ln(font_size * 0.3)  # Small space between paragraphs

        # Footer is now handled by the CustomPDF class

        temp_output = tempfile.NamedTemporaryFile(delete=False, suffix='.pdf')
        pdf.output(temp_output.name)
        return temp_output.name
        
    except Exception as e:
        logger.error(f"Error creating PDF: {str(e)}")
        # Fallback to basic PDF creation if advanced formatting fails
        return create_pdf_with_text_basic(text_content)

def create_docx_with_text(text_content, font_family='helvetica', font_size=12, page_size='A4',
                         orientation='portrait', margin=15, line_spacing=1.5, alignment='left',
                         include_page_numbers=True, header_text='', footer_text=''):
    """Create a Word document with the given text content using the specified formatting options"""
    try:
        # Import python-docx library - will fail if not installed
        from docx import Document
        from docx.shared import Pt, Inches, Mm
        from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
        
        # Log settings for debugging
        logger.info(f"Creating DOCX with settings: font={font_family}, size={font_size}, page={page_size}, orientation={orientation}")
        
        # Create a new document
        doc = Document()
        
        # Set page size and orientation
        section = doc.sections[0]
        if page_size == 'A4':
            if orientation == 'portrait':
                section.page_width, section.page_height = Mm(210), Mm(297)
            else:
                section.page_width, section.page_height = Mm(297), Mm(210)
        elif page_size == 'Letter':
            if orientation == 'portrait':
                section.page_width, section.page_height = Inches(8.5), Inches(11)
            else:
                section.page_width, section.page_height = Inches(11), Inches(8.5)
        
        # Set margins
        section.left_margin = Mm(margin)
        section.right_margin = Mm(margin)
        section.top_margin = Mm(margin)
        section.bottom_margin = Mm(margin)
        
        # Add header if specified
        if header_text:
            header = section.header
            header_para = header.paragraphs[0]
            header_para.text = header_text
            header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            header_para.style.font.size = Pt(font_size - 2)
            header_para.style.font.italic = True
        
        # Add footer if specified
        if footer_text or include_page_numbers:
            footer = section.footer
            footer_para = footer.paragraphs[0]
            
            if footer_text and include_page_numbers:
                footer_para.text = f"{footer_text} | "
                footer_para.add_run().add_field('PAGE')
            elif footer_text:
                footer_para.text = footer_text
            elif include_page_numbers:
                footer_para.add_run().add_field('PAGE')
                
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            footer_para.style.font.size = Pt(font_size - 2)
            footer_para.style.font.italic = True
        
        # Set paragraph alignment
        align_dict = {
            'left': WD_ALIGN_PARAGRAPH.LEFT,
            'center': WD_ALIGN_PARAGRAPH.CENTER,
            'right': WD_ALIGN_PARAGRAPH.RIGHT,
            'justified': WD_ALIGN_PARAGRAPH.JUSTIFY
        }
        para_align = align_dict.get(alignment, WD_ALIGN_PARAGRAPH.LEFT)
        
        # Add text content
        paragraphs = text_content.split('\n')
        for paragraph in paragraphs:
            if not paragraph.strip():
                doc.add_paragraph()
                continue
                
            para = doc.add_paragraph(paragraph)
            para.alignment = para_align
            
            # Set font properties
            for run in para.runs:
                run.font.size = Pt(font_size)
                if font_family == 'helvetica':
                    run.font.name = 'Arial'
                elif font_family == 'times':
                    run.font.name = 'Times New Roman'
                elif font_family == 'courier':
                    run.font.name = 'Courier New'
            
            # Set line spacing
            if line_spacing == 1.0:
                para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            elif line_spacing == 1.5:
                para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            elif line_spacing == 2.0:
                para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
            else:
                para.paragraph_format.line_spacing = line_spacing
        
        # Save the document to a temporary file with proper error handling
        temp_output = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
        try:
            doc.save(temp_output.name)
            logger.info(f"DOCX created successfully at: {temp_output.name}")
            return temp_output.name
        except Exception as save_error:
            logger.error(f"Error saving DOCX file: {str(save_error)}")
            raise Exception(f"Could not save DOCX file: {str(save_error)}")
        
    except ImportError as import_error:
        logger.error(f"python-docx library not installed: {str(import_error)}")
        logger.info("Falling back to PDF format")
        # Fall back to PDF if python-docx is not available
        return create_pdf_with_formatting(
            text_content, font_family, font_size, page_size, 
            orientation, margin, line_spacing, alignment,
            include_page_numbers, header_text, footer_text
        )
    except Exception as e:
        logger.error(f"Error creating DOCX: {str(e)}")
        logger.info("Falling back to PDF format")
        # Also fall back to PDF on any other error
        return create_pdf_with_formatting(
            text_content, font_family, font_size, page_size, 
            orientation, margin, line_spacing, alignment,
            include_page_numbers, header_text, footer_text
        )

def create_html_with_text(text_content, font_family='helvetica', font_size=12, 
                         line_spacing=1.5, alignment='left',
                         include_page_numbers=False, header_text='', footer_text=''):
    """Create an HTML file with the given text content using the specified formatting options"""
    # Map font families to CSS equivalents
    font_map = {
        'helvetica': 'Arial, Helvetica, sans-serif',
        'times': 'Times New Roman, Times, serif',
        'courier': 'Courier New, Courier, monospace'
    }
    font_css = font_map.get(font_family, 'Arial, Helvetica, sans-serif')
    
    # Map alignment to CSS
    align_css = alignment
    
    # Create HTML content
    html_content = f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <title>Translated Document</title>
    <style>
        body {{
            font-family: {font_css};
            font-size: {font_size}pt;
            line-height: {line_spacing};
            text-align: {align_css};
            margin: 30px;
        }}
        .header {{
            text-align: center;
            font-style: italic;
            margin-bottom: 30px;
            font-size: {font_size - 2}pt;
        }}
        .footer {{
            text-align: center;
            font-style: italic;
            margin-top: 30px;
            font-size: {font_size - 2}pt;
        }}
        p {{
            margin-bottom: 15px;
        }}
    </style>
</head>
<body>
"""

    # Add header if specified
    if header_text:
        html_content += f'<div class="header">{header_text}</div>\n'
    
    # Add content
    paragraphs = text_content.split('\n')
    for paragraph in paragraphs:
        if paragraph.strip():
            html_content += f'<p>{paragraph}</p>\n'
        else:
            html_content += '<p>&nbsp;</p>\n'  # Empty paragraph
    
    # Add footer if specified
    if footer_text:
        html_content += f'<div class="footer">{footer_text}</div>\n'
    
    html_content += """
</body>
</html>
"""
    
    # Write to file
    try:
        temp_output = tempfile.NamedTemporaryFile(delete=False, suffix='.html')
        with open(temp_output.name, 'w', encoding='utf-8') as f:
            f.write(html_content)
        return temp_output.name
    except Exception as e:
        logger.error(f"Error creating HTML: {str(e)}")
        raise Exception(f"Failed to create HTML document: {str(e)}")

def analyze_complexity(text):
    """Analyze text complexity to determine if it needs AI review.
    
    Returns a complexity score and features dict with analysis details.
    Higher scores indicate more complex text that would benefit from review.
    """
    if not text or not isinstance(text, str):
        return 0, {}
        
    # Initialize complexity metrics
    features = {
        'length': len(text),
        'avg_sentence_length': 0,
        'long_sentences': 0,  # Sentences > 30 words
        'complex_words': 0,   # Words > 6 chars
        'punctuation_count': 0,
        'special_chars': 0
    }
    
    # Count punctuation and special characters
    punctuation_chars = '.,:;?!-—()[]{}"\''
    special_chars = '@#$%^&*+=<>|~`§±'
    
    for char in text:
        if char in punctuation_chars:
            features['punctuation_count'] += 1
        if char in special_chars:
            features['special_chars'] += 1
    
    # Split into sentences and words
    sentences = [s.strip() for s in text.replace('!', '.').replace('?', '.').split('.') if s.strip()]
    if sentences:
        features['avg_sentence_length'] = len(text) / len(sentences)
        
    # Count long sentences and complex words
    for sentence in sentences:
        words = sentence.split()
        if len(words) > 30:
            features['long_sentences'] += 1
        
        for word in words:
            if len(word) > 6:
                features['complex_words'] += 1
    
    # Calculate complexity score (weighted sum of features)
    complexity_score = 0
    
    # Base score on length (longer text is more likely to need review)
    if features['length'] > 5000:
        complexity_score += 30
    elif features['length'] > 1000:
        complexity_score += 20
    elif features['length'] > 500:
        complexity_score += 10
        
    # Score based on sentence structure
    if features['avg_sentence_length'] > 25:
        complexity_score += 25
    elif features['avg_sentence_length'] > 15:
        complexity_score += 15
        
    # Score based on complex words ratio
    if features['complex_words'] > 100:
        complexity_score += 20
    elif features['complex_words'] > 50:
        complexity_score += 10
        
    # Score based on punctuation density
    punctuation_ratio = features['punctuation_count'] / max(1, features['length'])
    if punctuation_ratio > 0.1:
        complexity_score += 15
    elif punctuation_ratio > 0.05:
        complexity_score += 10
        
    # Score special characters
    if features['special_chars'] > 20:
        complexity_score += 10
        
    # Add bonus for long sentences
    complexity_score += min(20, features['long_sentences'] * 5)
    
    # Cap at 100
    complexity_score = min(100, complexity_score)
    
    return complexity_score, features

def process_document(filepath, deepl_api_key, openai_api_key=None, assistant_id=None, source_language='auto', target_language='SV', custom_instructions=None, return_segments=False, use_cache=True, smart_review=False, complexity_threshold=40, glossary_id=None):
    """Process a document by extracting text, translating with DeepL.
    
    Works with various file formats including PDF, DOCX, DOC, TXT, RTF, and ODT.
    Uses translation caching to improve performance and reduce API calls.
    With glossary_id, applies custom glossary terms to translations.
    
    The OpenAI review step has been separated into an optional post-processing step.
    
    Returns a tuple containing:
    1. Either the processed pdf bytes, or the list of translation segments
    2. Stats dictionary with cache_hits, cache_ratio, glossary_hits, glossary_ratio, and unique_terms_used
    """
    try:
        # Extract text from the file using the appropriate method
        file_ext = os.path.splitext(filepath)[1].lower()
        try:
            text_sections = extract_text_from_file(filepath)
            total_sections = len(text_sections)
            logger.info(f"Starting to process document with {total_sections} sections")
        except Exception as e:
            logger.error(f"Failed to extract text from file: {str(e)}")
            raise Exception(f"Could not read document content: {str(e)}")
        
        logger.info(f"Translation settings - Source: {source_language}, Target: {target_language}")
        if glossary_id:
            logger.info(f"Using glossary ID: {glossary_id}")
            
        # Process each section
        translations = []
        cache_hits = 0
        glossary_applied_count = 0
        total_glossary_hits = 0
        unique_glossary_terms = set()
        
        for i, section in enumerate(text_sections):
            try:
                section_id = section['id']
                original_text = section['text']
                source_info = section.get('source', 'unknown')
                
                logger.info(f"Processing section {i+1}/{total_sections} (id: {section_id}, source: {source_info})")
                
                if not original_text or len(original_text.strip()) < 10:
                    logger.warning(f"Section {i+1} contains insufficient text, skipping")
                    continue
                
                # Step 1: Translate text using DeepL (with caching and glossary)
                logger.info(f"Translating section {i+1} with DeepL")
                text_hash = None
                source_text_for_cache = None
                section_glossary_hits = 0
                section_glossary_terms = 0
                
                try:
                    # The updated translate_text function returns an extended tuple
                    translation_result = translate_text(
                        original_text, 
                        deepl_api_key, 
                        target_language, 
                        source_language, 
                        use_cache=use_cache,
                        glossary_id=glossary_id
                    )
                    
                    # Handle tuple return
                    if isinstance(translation_result, tuple):
                        if len(translation_result) >= 5:  # New version with glossary stats
                            translated_text, text_hash, source_text_for_cache, section_glossary_hits, section_glossary_terms = translation_result
                        elif len(translation_result) == 3:  # Old version
                            translated_text, text_hash, source_text_for_cache = translation_result
                        else:
                            translated_text = translation_result[0]
                            
                        # If we got translation from cache, increment counter
                        if text_hash and source_text_for_cache == original_text:
                            cache_hits += 1
                    else:
                        translated_text = translation_result
                        
                    if not translated_text:
                        logger.error("Translation returned empty result")
                        translated_text = original_text
                        raise ValueError("Translation returned empty result")
                        
                    # Track glossary stats
                    if glossary_id:
                        glossary_applied_count += 1
                        total_glossary_hits += section_glossary_hits
                        # The unique terms set will accumulate across all sections
                        if section_glossary_terms > 0:
                            # Since we don't have the actual term IDs from this level, 
                            # we'll use a placeholder to track this section's unique terms count
                            for i in range(section_glossary_terms):
                                unique_glossary_terms.add(f"section_{section_id}_term_{i}")
                except Exception as e:
                    logger.error(f"DeepL translation error for section {i+1}: {str(e)}")
                    translated_text = original_text  # Fallback to original text
                
                # Prepare metadata for saving to cache (if this wasn't from cache already)
                cache_metadata = {}
                if use_cache and text_hash and source_text_for_cache:
                    cache_metadata = {
                        'source_text': source_text_for_cache,
                        'source_hash': text_hash,
                        'target_language': target_language
                    }
                
                # Calculate complexity score for information purposes
                complexity_score = 0
                complexity_features = {}
                if openai_api_key and assistant_id:
                    complexity_score, complexity_features = analyze_complexity(original_text)
                    logger.info(f"Section {i+1} complexity score: {complexity_score}")
                
                # Extra check: ensure translated_text is different from original_text
                # If they're identical, log a warning and set translated_text to empty
                if translated_text and translated_text == original_text and source_language != target_language:
                    logger.warning(f"Translation for section {section_id} is identical to original. This may indicate an issue with the translation service.")
                    logger.warning(f"Source language: {source_language}, Target language: {target_language}")
                    logger.warning(f"Setting translated text to empty for now - user can edit it manually")
                    translated_text = ""
                
                translations.append({
                    'id': section_id,
                    'original_text': original_text,
                    'translated_text': translated_text,
                    'status': 'success',
                    'source': source_info,
                    'cache_metadata': cache_metadata,
                    'complexity_score': complexity_score,
                    'reviewed_by_ai': False,  # AI review is now a separate step
                    'glossary_applied': glossary_id is not None,
                    'glossary_hits': section_glossary_hits,
                    'glossary_terms': section_glossary_terms
                })
                logger.info(f"Successfully completed processing section {i+1}")
                
            except Exception as e:
                logger.error(f"Error processing section {i+1}: {str(e)}")
                translations.append({
                    'id': section['id'] if 'section' in locals() else i,
                    'original_text': original_text if 'original_text' in locals() else '',
                    'translated_text': translated_text if 'translated_text' in locals() else original_text if 'original_text' in locals() else '',
                    'status': 'error',
                    'error': str(e),
                    'source': section.get('source', 'unknown') if 'section' in locals() else 'unknown'
                })
                continue
        
        # Calculate statistics
        stats = {}
        if total_sections > 0:
            # Cache statistics
            if use_cache:
                cache_rate = (cache_hits / total_sections) * 100
                stats['cache_hits'] = cache_hits
                stats['cache_ratio'] = cache_rate
                logger.info(f"Translation cache performance: {cache_hits}/{total_sections} sections from cache ({cache_rate:.1f}%)")
            
            # Glossary statistics
            if glossary_id:
                glossary_rate = (glossary_applied_count / total_sections) * 100
                stats['glossary_hits'] = total_glossary_hits
                stats['glossary_ratio'] = (total_glossary_hits / len(''.join([t.get('original_text', '') for t in translations]))) * 1000  # Per 1000 characters
                stats['unique_terms_used'] = len(unique_glossary_terms)
                
                logger.info(f"Glossary applied to {glossary_applied_count}/{total_sections} sections ({glossary_rate:.1f}%)")
                logger.info(f"Total glossary hits: {total_glossary_hits} terms replaced")
                logger.info(f"Unique glossary terms used: {len(unique_glossary_terms)}")
        
        # Sort translations by section ID to maintain document order
        translations.sort(key=lambda x: x['id'])
        
        successful_sections = sum(1 for t in translations if t['status'] == 'success')
        logger.info(f"Document processing complete. Successfully processed {successful_sections} out of {total_sections} sections")
        
        if not translations:
            raise Exception("No sections could be processed")
        
        if return_segments:
            return translations, stats
        else:
            successful_translations = [t['translated_text'] for t in translations if t['status'] == 'success']
            if not successful_translations:
                raise Exception("No sections were successfully translated")
            combined_text = '\n\n'.join(successful_translations)
            return create_pdf_with_text(combined_text), stats
    
    except Exception as e:
        logger.error(f"Document processing error: {str(e)}")
        raise Exception(f"Document processing failed: {str(e)}")

def review_page_translation(page_content, openai_api_key, assistant_id, custom_instructions=None):
    """Review a single page translation using OpenAI assistant.
    
    This function sends the translated text to OpenAI for review and returns
    the reviewed translation. It is meant to be called on a page-by-page basis.
    
    Args:
        page_content: The translated text to review
        openai_api_key: OpenAI API key
        assistant_id: ID of the assistant to use
        custom_instructions: Custom instructions for the review (optional)
        
    Returns:
        tuple: (reviewed_text, success_flag, error_message)
            - reviewed_text: The text after AI review, or original if review failed
            - success_flag: True if review was successful, False otherwise
            - error_message: Error message if review failed, None otherwise
    """
    if not page_content or not openai_api_key or not assistant_id:
        return page_content, False, "Missing required parameters"
    
    try:
        logger.info(f"Reviewing page translation with OpenAI (content length: {len(page_content)})")
        
        # Calculate complexity score for information
        complexity_score, _ = analyze_complexity(page_content)
        logger.info(f"Page complexity score: {complexity_score}")
        
        # Send to review_translation
        reviewed_text = review_translation(
            page_content,
            openai_api_key,
            assistant_id,
            instructions=custom_instructions
        )
        
        if reviewed_text and reviewed_text != page_content:
            logger.info("Successfully reviewed page translation")
            return reviewed_text, True, None
        elif reviewed_text:
            logger.info("Review returned unchanged text")
            return reviewed_text, True, None
        else:
            logger.warning("Review returned empty text")
            return page_content, False, "Review returned empty text"
            
    except Exception as e:
        logger.error(f"Error reviewing page translation: {str(e)}")
        return page_content, False, str(e)

# For backward compatibility
def process_pdf(filepath, deepl_api_key, openai_api_key, assistant_id, source_language='auto', target_language='SV', custom_instructions=None, return_segments=False):
    """Legacy function for backward compatibility"""
    return process_document(filepath, deepl_api_key, openai_api_key, assistant_id, source_language, target_language, custom_instructions, return_segments)